# Full, final Streamlit app (core parts). Merge into your file replacing the previous implementations.
# Note: Keep your original imports (streamlit, pandas, fitz, huggingface_hub.InferenceClient, matplotlib, etc.)
# This block focuses on the improved logic. I've kept style similar to your earlier code.

import os
import re
import time
import streamlit as st
import io
import pandas as pd
from difflib import get_close_matches
from huggingface_hub import InferenceClient
from docx import Document
import matplotlib.pyplot as plt
from typing import Dict, List, Tuple

UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)
HISTORY_FILE = "qa_history.txt"

# ---- Helper: file extraction (unchanged from your working code) ----
import fitz  # PyMuPDF

def extract_pdf_text(file_path):
    text = ""
    with fitz.open(file_path) as doc:
        for page in doc:
            text += page.get_text()
    return text

def extract_excel_data(file_path):
    dfs = pd.read_excel(file_path, sheet_name=None)
    return dfs

def extract_docx_text(file_path):
    doc = Document(file_path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        table_text = []
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            table_text.append("\t".join(row_text))
        if table_text:
            parts.append("\n".join(table_text))
    return "\n".join(parts)

def load_files():
    """
    Load uploaded files from UPLOAD_DIR.
    Returns:
      file_data: dict filename -> text/summary
      excel_files: dict filename -> dict(sheet_name -> DataFrame)
    """
    file_data = {}
    excel_files = {}
    for fname in os.listdir(UPLOAD_DIR):
        path = os.path.join(UPLOAD_DIR, fname)
        if fname.lower().endswith(".pdf"):
            file_data[fname] = extract_pdf_text(path)
        elif fname.lower().endswith((".xls", ".xlsx")):
            dfs = extract_excel_data(path)
            excel_files[fname] = dfs
            file_data[fname] = f"Excel file with {len(dfs)} sheets loaded."
        elif fname.lower().endswith(".docx"):
            file_data[fname] = extract_docx_text(path)
    return file_data, excel_files

# ---- Hugging Face query helper (token conservative) ----
HF_MAX_TOKENS = 120  # keep low for free tier

def query_model(client, model, user_prompt, max_tokens=HF_MAX_TOKENS):
    """Try chat completion, fallback to text generation. Keep tokens small."""
    try:
        messages = [{"role": "user", "content": user_prompt}]
        resp = client.chat_completion(model=model, messages=messages, max_tokens=max_tokens)
        return resp.choices[0].message["content"]
    except Exception as e1:
        # fallback to text_generation
        try:
            resp = client.text_generation(model=model, prompt=user_prompt, max_new_tokens=max_tokens)
            if isinstance(resp, str):
                return resp
            return resp.get("generated_text", str(resp))
        except Exception as e2:
            return f"Model calls failed: {e1} | {e2}"

def append_to_history_file(question, answer):
    with open(HISTORY_FILE, "a", encoding="utf-8") as f:
        f.write(f"Q: {question}\nA: {answer}\n\n")

# ---- Column mapping + synonyms ----
def normalize_columns(columns):
    """
    columns: iterable -> normalized_key -> original
    normalized uses lower-case + underscores
    """
    mapping = {}
    for col in columns:
        key = re.sub(r'\s+', '_', str(col).lower()).strip()
        mapping[key] = col
    return mapping

# A small synonyms dictionary you can extend
DEFAULT_SYNONYMS = {
    "doctor": "surgeon",
    "doc": "surgeon",
    "manager": "reporting manager",
    "mgr": "reporting manager",
    "price": "product_1_price",
    "revenue": "product_1_price",
    "amount": "product_1_price",
    "surgeries": "surgery_no",
    "surgery": "surgery_no",
    "hospital": "hospital",
    "state": "state",
    "zone": "zone",
}

def map_query_to_columns(query: str, df_columns: List[str], synonyms: Dict[str,str]=None) -> List[str]:
    """
    Map tokens/phrases in the query to actual DataFrame column names.
    Returns ordered unique list of mapped original column names.
    """
    if synonyms is None:
        synonyms = DEFAULT_SYNONYMS

    normalized = normalize_columns(df_columns)  # norm_key -> original
    q = query.lower()

    mapped = []

    # First: check for full phrase matches (prefer longer names)
    # Sort normalized keys by length desc to match multi-word columns first
    for nkey in sorted(normalized.keys(), key=len, reverse=True):
        if nkey in re.sub(r'\s+', '_', q):
            mapped.append(normalized[nkey])

    # Then token-level matching + synonyms + fuzzy
    tokens = re.findall(r'\w+', q)
    for token in tokens:
        # synonyms
        if token in synonyms:
            syn = synonyms[token]
            syn_key = re.sub(r'\s+','_', syn.lower())
            if syn_key in normalized:
                mapped.append(normalized[syn_key])
                continue
        # direct normalized token match
        if token in normalized:
            mapped.append(normalized[token])
            continue
        # fuzzy match on normalized keys
        close = get_close_matches(token, normalized.keys(), n=1, cutoff=0.78)
        if close:
            mapped.append(normalized[close[0]])

    # Deduplicate preserving order
    result = []
    for c in mapped:
        if c not in result:
            result.append(c)
    return result

# ---- The full robust query function (final version) ----
def query_excel_data(sheets_dict: Dict[str, pd.DataFrame], question: str) -> List[Tuple[str, pd.DataFrame]]:
    """
    Query Excel sheets (sheets_dict: {sheet_name: DataFrame}) using natural language question.
    Returns list of (sheet_name, DataFrame) results (each DataFrame is truncated for display).
    """
    q = question.strip()
    q_lower = q.lower()
    results = []

    # Pre-parse top-level query intent tokens
    wants_count = bool(re.search(r'\b(count|how many|number of|total rows)\b', q_lower))
    wants_sum = bool(re.search(r'\b(sum|total|revenue|amount)\b', q_lower))
    wants_avg = bool(re.search(r'\b(average|avg|mean)\b', q_lower))
    wants_max = bool(re.search(r'\b(max|maximum|highest)\b', q_lower))
    wants_min = bool(re.search(r'\b(min|minimum|lowest)\b', q_lower))
    wants_trend = bool(re.search(r'\b(trend|over time|over the years|by year|by quarter|by month)\b', q_lower))
    top_bottom_match = re.search(r'\b(top|bottom)\s+(\d+)\b', q_lower)

    # Parse filter conditions: supports multiple conditions like "state = karnataka and year = 2024"
    # It finds occurrences of "<col> <op> <value>"
    cond_pattern = re.findall(r'([\w\s]+?)\s*(>=|<=|=|>|<)\s*([^\s,;]+)', q, flags=re.IGNORECASE)

    # Helper: parse group by
    by_match = re.search(r'\bby\s+([\w\s_]+)', q_lower)
    group_by_col_name = None
    if by_match:
        group_by_col_name = by_match.group(1).strip()

    for sheet_name, df in sheets_dict.items():
        if df is None or df.empty:
            continue

        df_copy = df.copy()
        # Normalize column names to strings
        df_copy.columns = [str(c).strip() for c in df_copy.columns]
        # Normalize string columns to trimmed lowercase for reliable matching
        str_cols = df_copy.select_dtypes(include=["object", "string"]).columns.tolist()
        for c in str_cols:
            # keep original dtype for non-string later by using df_copy_for_ops if need be
            df_copy[c] = df_copy[c].astype(str).str.strip().str.lower()

        numeric_cols = df_copy.select_dtypes(include="number").columns.tolist()

        # ---------- Apply Filters ----------
        for raw_col_name, op, raw_val in cond_pattern:
            # map query fragment to real column
            mapped_cols = map_query_to_columns(raw_col_name, df_copy.columns)
            if not mapped_cols:
                continue
            col = mapped_cols[0]  # choose best match

            # Try numeric compare first if column is numeric (note: df_copy[col] may be str if earlier converted; rely on original df)
            if col in numeric_cols:
                try:
                    val_num = float(raw_val)
                except:
                    # can't convert => skip this condition
                    continue
                if op == "=":
                    df_copy = df_copy[df_copy[col] == val_num]
                elif op == ">":
                    df_copy = df_copy[df_copy[col] > val_num]
                elif op == "<":
                    df_copy = df_copy[df_copy[col] < val_num]
                elif op == ">=":
                    df_copy = df_copy[df_copy[col] >= val_num]
                elif op == "<=":
                    df_copy = df_copy[df_copy[col] <= val_num]
            else:
                # string column - do strict equality for '=' operator
                val_norm = str(raw_val).strip().lower()
                col_series = df_copy[col].astype(str).str.strip().str.lower()
                if op == "=":
                    # strict equality first
                    matched = (col_series == val_norm)
                    if matched.any():
                        df_copy = df_copy[matched]
                    else:
                        # fallback: fuzzy match (catch small typos)
                        unique_vals = pd.Series(col_series.unique()).dropna().astype(str).values.tolist()
                        close = get_close_matches(val_norm, unique_vals, n=1, cutoff=0.86)
                        if close:
                            df_copy = df_copy[col_series == close[0]]
                        else:
                            # no match -> filter to nothing
                            df_copy = df_copy[col_series == val_norm]  # empty
                else:
                    # other ops on text: treat as contains (case-insensitive)
                    df_copy = df_copy[col_series.str.contains(val_norm, na=False)]

        # If after filtering dataframe is empty, return empty result for this sheet
        if df_copy.empty:
            results.append((sheet_name, df_copy.head(0)))
            continue

        # ---------- Determine columns to project (if any) ----------
        # If user mentioned column names in query (e.g., "surgeon reporting manager designation"), pick them
        projection_cols = map_query_to_columns(q, df_copy.columns)
        # Remove columns that are part of "group by" phrase to avoid confusion
        if group_by_col_name:
            mapped_group_by = map_query_to_columns(group_by_col_name, df_copy.columns)
            if mapped_group_by and mapped_group_by[0] in projection_cols:
                projection_cols.remove(mapped_group_by[0])

        # ---------- Aggregations / Group-by ----------
        agg_result = None
        # Group-by column if asked and valid
        group_by_col = None
        if group_by_col_name:
            mapped_g = map_query_to_columns(group_by_col_name, df_copy.columns)
            if mapped_g:
                group_by_col = mapped_g[0]

        # Helper to pick a target numeric column from query or fallback
        def pick_target_numeric():
            cand = map_query_to_columns(q, df_copy.columns)
            for c in cand:
                if c in numeric_cols:
                    return c
            return numeric_cols[0] if numeric_cols else None

        # count
        if wants_count and not (wants_sum or wants_avg or wants_max or wants_min):
            if group_by_col:
                agg_result = df_copy.groupby(group_by_col).size().reset_index(name="Count")
            else:
                agg_result = pd.DataFrame({"Count": [len(df_copy)]})

        # sum / average / min / max with optional group-by
        elif wants_sum or wants_avg or wants_max or wants_min:
            target = pick_target_numeric()
            if target:
                if group_by_col:
                    if wants_sum:
                        agg_result = df_copy.groupby(group_by_col, dropna=False)[target].sum().reset_index()
                    elif wants_avg:
                        agg_result = df_copy.groupby(group_by_col, dropna=False)[target].mean().reset_index()
                    elif wants_max:
                        agg_result = df_copy.groupby(group_by_col, dropna=False)[target].max().reset_index()
                    elif wants_min:
                        agg_result = df_copy.groupby(group_by_col, dropna=False)[target].min().reset_index()
                else:
                    if wants_sum:
                        agg_result = pd.DataFrame({target: [df_copy[target].sum()]})
                    elif wants_avg:
                        agg_result = pd.DataFrame({target: [df_copy[target].mean()]})
                    elif wants_max:
                        agg_result = pd.DataFrame({target: [df_copy[target].max()]})
                    elif wants_min:
                        agg_result = pd.DataFrame({target: [df_copy[target].min()]})

        # top/bottom
        if top_bottom_match:
            tb = top_bottom_match.group(1)
            n = int(top_bottom_match.group(2))
            col_for_sort = pick_target_numeric()
            if col_for_sort:
                if tb == "top":
                    agg_result = df_copy.nlargest(n, col_for_sort)
                else:
                    agg_result = df_copy.nsmallest(n, col_for_sort)

        # trend
        if wants_trend and agg_result is None:
            # try to find year/month/quarter column
            for time_col in ["year", "month", "quarter"]:
                mapped_t = map_query_to_columns(time_col, df_copy.columns)
                if mapped_t:
                    tcol = mapped_t[0]
                    if numeric_cols:
                        agg_result = df_copy.groupby(tcol)[numeric_cols].sum().reset_index()
                        break

        # ---------- Final selection & output ----------
        if agg_result is not None:
            # truncate for display, but keep full for download if needed
            display_df = agg_result.head(200)
            results.append((sheet_name, display_df))
        else:
            # No aggregation requested -> return filtered table with projection if requested
            if projection_cols:
                # ensure selected columns exist
                cols = [c for c in projection_cols if c in df_copy.columns]
                if not cols:
                    # fallback show head
                    results.append((sheet_name, df_copy.head(200)))
                else:
                    # de-duplicate rows so Surgeon list doesn't repeat many times
                    display_df = df_copy[cols].drop_duplicates().head(200)
                    results.append((sheet_name, display_df))
            else:
                # default: return first N rows of filtered dataset
                results.append((sheet_name, df_copy.head(200)))

    return results if results else None

# ---- small plotting helper (unchanged, but uses normalized df) ----
def try_plot(question, excel_sheets_or_df, answer=None):
    keywords = ["compare", "trend", "graph", "plot", "visualize", "chart", "over time"]
    if not any(k in question.lower() for k in keywords):
        return None

    if isinstance(excel_sheets_or_df, dict):
        sheets_iter = excel_sheets_or_df.items()
    elif isinstance(excel_sheets_or_df, pd.DataFrame):
        sheets_iter = [("sheet", excel_sheets_or_df)]
    else:
        sheets_iter = []

    for sheet, df in sheets_iter:
        if df is None or df.empty:
            continue
        # If a Year-like column exists and numeric columns available
        possible_year_cols = [c for c in df.columns if c.lower() in ("year", "month", "quarter")]
        numeric_cols = df.select_dtypes(include="number").columns.tolist()
        if possible_year_cols and numeric_cols:
            xcol = possible_year_cols[0]
            ycols = [c for c in numeric_cols if c != xcol]
            if not ycols:
                ycols = numeric_cols
            if not ycols:
                continue
            fig, ax = plt.subplots()
            df.plot(x=xcol, y=ycols, kind="line", ax=ax, marker="o")
            ax.set_title(f"Trend from {sheet}")
            st.pyplot(fig)
            return f"Visualization generated from sheet '{sheet}'"

    if answer:
        # try extracting simple year-revenue style table
        lines = answer.splitlines()
        data = []
        for line in lines:
            m = re.match(r"^(\d{4})\s*[:\-]?\s*[€$]?([\d,\.]+)", line.strip())
            if m:
                data.append([m.group(1), float(m.group(2).replace(",", ""))])
        if data:
            dfp = pd.DataFrame(data, columns=["Year", "Revenue"])
            fig, ax = plt.subplots()
            dfp.plot(x="Year", y="Revenue", kind="bar", ax=ax)
            st.pyplot(fig)
            return "Visualization generated from model's answer"

    return "No suitable numeric data found for visualization."

# ---------------------
# Streamlit UI glue (kept similar to your original)
# ---------------------
st.set_page_config(page_title="Local File Chatbot", layout="wide")
st.title("Local File Chatbot")

# Sidebar: HF API key + model selection
st.sidebar.header("Hugging Face Setup")
api_key = st.sidebar.text_input("Enter your Hugging Face API Key:", type="password")
model_choice = st.sidebar.selectbox(
    "Select Model",
    ["deepseek-ai/DeepSeek-V3", "HuggingFaceH4/zephyr-7b-beta", "google/flan-t5-base"],
    index=0
)

client = None
if api_key:
    try:
        client = InferenceClient(api_key=api_key, provider="auto")
        st.sidebar.success("API Key configured")
    except Exception as e:
        st.sidebar.error(f"API client error: {e}")
        client = None
else:
    st.sidebar.warning("Hugging Face API key not set. Model features disabled.")

# File uploader (unchanged)
uploaded_file = st.file_uploader("Upload PDF, Excel, or Word", type=["pdf", "xls", "xlsx", "docx"])
if uploaded_file is not None:
    file_path = os.path.join(UPLOAD_DIR, uploaded_file.name)
    with st.spinner(f"Saving {uploaded_file.name}..."):
        time.sleep(0.4)
        with open(file_path, "wb") as f:
            f.write(uploaded_file.getbuffer())
    st.success(f"Saved {uploaded_file.name}")

# Load files
files_content, excel_files = load_files()

if files_content:
    st.subheader("Loaded Files")
    selected_file = st.selectbox("Select a file to ask questions about:", list(files_content.keys()))
    selected_file_content = files_content[selected_file]
else:
    st.info("No files uploaded yet. Please upload a PDF/Excel/Word file.")
    selected_file = None
    selected_file_content = ""

# Sheet selection when Excel selected
selected_sheet = None
if selected_file and selected_file.lower().endswith((".xls", ".xlsx")):
    sheets = list(excel_files.get(selected_file, {}).keys())
    if sheets:
        selected_sheet = st.selectbox("Select sheet (or choose (all)):", ["(all)"] + sheets)
    else:
        st.info("Excel file has no sheets detected.")

# Conversation state
if "history" not in st.session_state:
    st.session_state["history"] = []

# Helper: query the selected excel file/sheet
def query_excel_for_user(selected_file, selected_sheet, question):
    if not selected_file or selected_file not in excel_files:
        return None
    sheets_dict = excel_files[selected_file]
    if selected_sheet and selected_sheet != "(all)":
        # return only the chosen sheet
        return query_excel_data({selected_sheet: sheets_dict[selected_sheet]}, question)
    else:
        return query_excel_data(sheets_dict, question)

# Input area
user_question = st.text_input("Ask a question about your files:")
if st.button("Ask"):
    if not files_content:
        st.error("No files uploaded.")
    elif not user_question.strip():
        st.warning("Type a question first.")
    else:
        # If Excel selected, query that first
        results = None
        if selected_file and selected_file.lower().endswith((".xls", ".xlsx")):
            results = query_excel_for_user(selected_file, selected_sheet, user_question)

        if results:
            st.subheader("Structured Result from Excel")
            for sheet, res_df in results:
                st.markdown(f"**Sheet: {sheet}**")
                if res_df is None or res_df.empty:
                    st.write("No matching rows.")
                else:
                    st.dataframe(res_df)
                    # allow download of the shown dataframe
                    csv_bytes = res_df.to_csv(index=False).encode("utf-8")
                    st.download_button(
                        label=f"Download result from {sheet}",
                        data=csv_bytes,
                        file_name=f"result_{sheet}.csv",
                        mime="text/csv"
                    )
                    output = io.BytesIO()
                    with pd.ExcelWriter(output, engine="openpyxl") as writer:
                        res_df.to_excel(writer, index=False, sheet_name=sheet)
                    st.download_button(
                        label=f"Download result from {sheet} (Excel)",
                        data=output.getvalue(),
                        file_name=f"result_{sheet}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    )

            # Run a small model summary only if the HF client exists and user wants a concise summary
            if client:
                # build a tiny sample to send (aggregated result or first 10 rows)
                sample_df = results[0][1].head(10)
                sample_md = sample_df.to_markdown(index=False)
                prompt = (
                    f"Question: {user_question}\n\n"
                    "Here is a small sample of relevant data (already filtered/aggregated):\n"
                    f"{sample_md}\n\n"
                    "Please provide a concise answer in 1-2 short sentences."
                )
                with st.spinner("Summarizing (model)..."):
                    answer = query_model(client, model_choice, prompt, max_tokens=100)
                st.subheader("Answer Summary")
                st.write(answer)
                st.session_state["history"].append((user_question, answer))
                append_to_history_file(user_question, answer)
            else:
                st.info("Model not available (API key not configured) — displayed structured result only.")
                st.session_state["history"].append((user_question, "Structured result returned (no model)."))
        else:
            # Not Excel (or no structured result) -> use text of selected file (PDF/DOCX)
            if selected_file_content and client:
                prompt = (
                    f"Context (extracted text) from file '{selected_file}':\n"
                    f"{selected_file_content[:4000]}\n\n"  # only send first chunk to save tokens
                    f"User question: {user_question}\n\nAnswer concisely:"
                )
                with st.spinner("Querying model..."):
                    answer = query_model(client, model_choice, prompt, max_tokens=140)
                st.write(answer)
                st.session_state["history"].append((user_question, answer))
                append_to_history_file(user_question, answer)
            elif selected_file_content:
                # no model but we have a text file - show a helpful note
                st.error("No Hugging Face API key configured. Can't answer text-based questions without a model.")
            else:
                st.error("No content available to answer your question.")

# Conversation history display
st.subheader("Conversation History")
if st.session_state["history"]:
    for idx, (q, a) in enumerate(st.session_state["history"], 1):
        st.markdown(f"**Q{idx}:** {q}")
        st.markdown(f"**A{idx}:** {a}")
        st.markdown("---")

# Q&A download
if os.path.exists(HISTORY_FILE):
    with open(HISTORY_FILE, "r", encoding="utf-8") as f:
        history_text = f.read()
    st.sidebar.download_button(
        label="Download Q&A History",
        data=history_text,
        file_name="qa_history.txt",
        mime="text/plain"
    )
